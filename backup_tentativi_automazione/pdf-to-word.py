import os
import sys
from pathlib import Path
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configurazione
PDF_FOLDER = r"C:\Users\armon\Downloads\ilovepdf_extracted-pages"
OUTPUT_WORD = r"D:\Karaoke Manager\CATALOG_COMPLETE.docx"

def merge_pdfs():
    """Unisce tutti i 30 PDF in uno solo"""
    print("UNIONE DEI 30 PDF IN UN UNICO FILE...")
    print("=" * 70)

    merger = PdfWriter()

    for i in range(1, 31):
        filename = f"Lista basi Song Service Ottobre 2025-{i}.pdf"
        filepath = os.path.join(PDF_FOLDER, filename)

        if os.path.exists(filepath):
            print(f"[{i}/30] Aggiungendo: {filename}")
            reader = PdfReader(filepath)
            for page in reader.pages:
                merger.add_page(page)
        else:
            print(f"ATTENZIONE: File non trovato: {filename}")

    # Salva il PDF unito temporaneamente
    temp_pdf = os.path.join(PDF_FOLDER, "catalogo_completo_temp.pdf")
    with open(temp_pdf, "wb") as output_file:
        merger.write(output_file)

    print(f"\nPDF unito salvato temporaneamente in: {temp_pdf}\n")
    return temp_pdf

def pdf_to_word_simple(pdf_path):
    """Converte PDF in Word estraendo solo il testo"""
    print("CONVERSIONE PDF -> WORD...")
    print("=" * 70)

    # Crea documento Word
    doc = Document()

    # Leggi PDF
    reader = PdfReader(pdf_path)
    total_pages = len(reader.pages)

    print(f"Pagine totali: {total_pages}\n")

    for page_num, page in enumerate(reader.pages, 1):
        print(f"[{page_num}/{total_pages}] Processando pagina {page_num}...")

        # Estrai testo
        text = page.extract_text()

        if text:
            lines = text.split('\n')

            for line in lines:
                trimmed = line.strip()

                # Skip righe vuote o header/footer
                if not trimmed or len(trimmed) < 2:
                    continue

                if ('M-LIVE' in trimmed or
                    'CATALOGO COMPLETO' in trimmed or
                    'www.m-live.com' in trimmed or
                    'Pagina' in trimmed or
                    'Copyright' in trimmed):
                    continue

                # Aggiungi paragrafo al documento
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(trimmed)

                # Se è tutto maiuscolo (probabile cantante), rendi grassetto
                if trimmed == trimmed.upper() and any(c.isalpha() for c in trimmed):
                    if not trimmed.isdigit() and len(trimmed) > 2:
                        run.bold = True
                        run.font.size = Pt(12)
                else:
                    run.font.size = Pt(10)

    # Salva documento
    doc.save(OUTPUT_WORD)
    print(f"\nFile Word salvato in: {OUTPUT_WORD}")
    print(f"Dimensione: {os.path.getsize(OUTPUT_WORD) / 1024:.2f} KB")

    return OUTPUT_WORD

def main():
    print("CONVERSIONE PDF -> WORD\n")
    print("Cartella PDF:", PDF_FOLDER)
    print("Output Word:", OUTPUT_WORD)
    print("\n" + "=" * 70 + "\n")

    try:
        # Step 1: Unisci tutti i PDF
        merged_pdf = merge_pdfs()

        # Step 2: Converti in Word
        word_file = pdf_to_word_simple(merged_pdf)

        # Step 3: Pulisci file temporaneo
        if os.path.exists(merged_pdf):
            os.remove(merged_pdf)
            print(f"File temporaneo rimosso: {merged_pdf}")

        print("\n" + "=" * 70)
        print("CONVERSIONE COMPLETATA!")
        print("=" * 70)
        print(f"\nApri il file Word: {word_file}")
        print("Ora puoi processarlo con Node.js per estrarre cantanti e canzoni.")

    except Exception as e:
        print(f"\nERRORE: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
